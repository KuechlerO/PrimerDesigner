import io
from dataclasses import replace
from datetime import datetime
from typing import List, Optional, Sequence, Union

from primer_designer_app.models import PrimerSettingsModel, DesignResultsSummary
from primer_designer_app.utils.variant_info import (
    AllelicVariantInfo,
    IndelType,
    TranscriptVariantInfo,
)
from primer_designer_app.utils.display_utils import (
    REPORT_DISPLAY_FLANK,
    compute_report_display_bounds,
    shift_template_hits_for_display,
)
from primer_designer_app.utils.hgvs_display import (
    hgvs_input_on_plain,
    normalize_indel_type,
    template_bases_consumed_by_bracket,
)
from primer_designer_app.utils.helpers import create_hgvs_notation
from primer_designer_app.utils.amplicon_display import (
    amplicon_chrom_label,
    extract_amplicon_summary,
    format_penalty_score,
    truncate_product_seq,
)
from primer_designer_app.utils.insilico_analysis import insilico_reference_description
from primer_designer_app.utils.primer_utils import INSILICO_OK, PrimerPairResult
from primer_designer_app.utils.sv_storage import SV_WINDOW_ORDER

from django.conf import settings
from django.urls import reverse

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import logging

logger = logging.getLogger(__name__)

# Word highlight indices (docx.enum.text.WD_COLOR_INDEX values)
HIGHLIGHT_PRIMER = 3  # TURQUOISE
HIGHLIGHT_VARIANT = 4  # BRIGHT_GREEN
HIGHLIGHT_VCF = "vcf"  # custom light lavender (matches web .highlight-vcf)
HIGHLIGHT_SNP = 7  # YELLOW
HIGHLIGHT_SNP_CONFLICT = "orange"  # custom run shading (no Word preset)

HighlightColor = Union[int, str]

_CUSTOM_HIGHLIGHT_FILLS: dict[str, str] = {
    HIGHLIGHT_VCF: "E0D1FA",
    HIGHLIGHT_SNP_CONFLICT: "FFC761",
}


def _set_run_highlight(run, color: Optional[HighlightColor]) -> None:
    if color is None:
        return
    fill = _CUSTOM_HIGHLIGHT_FILLS.get(color) if isinstance(color, str) else None
    if fill:
        r_pr = run._element.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        r_pr.append(shd)
        return
    run.font.highlight_color = color


def _intervals_overlap(a0: int, a1: int, b0: int, b1: int) -> bool:
    return a0 <= b1 and b0 <= a1


def _apply_highlight_range(
    lookup: List[Optional[HighlightColor]], start: int, end: int, color: HighlightColor
) -> None:
    for i in range(max(0, start), min(len(lookup), end + 1)):
        lookup[i] = color


def _snp_binding_conflict(
    hit: dict,
    primer_f_start: int,
    primer_f_end: int,
    primer_r_start: int,
    primer_r_end: int,
) -> bool:
    ts = int(hit["template_start"])
    te = int(hit["template_end"])
    in_forward = ts <= primer_f_end and te >= primer_f_start
    in_reverse = ts <= primer_r_end and te >= primer_r_start
    return in_forward or in_reverse


def build_template_highlight_lookup(
    plain_len: int,
    *,
    vcf_hits: Sequence[dict],
    snp_hits: Sequence[dict],
    primer_pair: PrimerPairResult,
    skip_snp_interval: Optional[tuple[int, int]] = None,
) -> List[Optional[HighlightColor]]:
    """
    Map template indices to Word highlight colors for report sequence snippets.

    Layer order mirrors the web UI: VCF → SNP → primer → SNP binding conflict.
    Variant bracket styling is applied separately in visualize_sequence_as_docx.
    """
    lookup: List[Optional[HighlightColor]] = [None] * max(0, plain_len)

    for hit in vcf_hits or []:
        _apply_highlight_range(
            lookup,
            int(hit["template_start"]),
            int(hit["template_end"]),
            HIGHLIGHT_VCF,
        )

    for hit in snp_hits or []:
        ts = int(hit["template_start"])
        te = int(hit["template_end"])
        if skip_snp_interval and _intervals_overlap(
            ts, te, skip_snp_interval[0], skip_snp_interval[1]
        ):
            continue
        _apply_highlight_range(lookup, ts, te, HIGHLIGHT_SNP)

    pf_s = int(primer_pair.left_relPos_start)
    pf_e = int(primer_pair.left_relPos_end)
    pr_s = int(primer_pair.right_relPos_start)
    pr_e = int(primer_pair.right_relPos_end)
    _apply_highlight_range(lookup, pf_s, pf_e, HIGHLIGHT_PRIMER)
    _apply_highlight_range(lookup, pr_s, pr_e, HIGHLIGHT_PRIMER)

    for hit in snp_hits or []:
        ts = int(hit["template_start"])
        te = int(hit["template_end"])
        if skip_snp_interval and _intervals_overlap(
            ts, te, skip_snp_interval[0], skip_snp_interval[1]
        ):
            continue
        if _snp_binding_conflict(hit, pf_s, pf_e, pr_s, pr_e):
            _apply_highlight_range(lookup, ts, te, HIGHLIGHT_SNP_CONFLICT)

    return lookup


def _report_region_hits(
    design_summary: DesignResultsSummary,
    var_info: AllelicVariantInfo,
) -> tuple[list[dict], list[dict]]:
    """VCF and gnomAD SNP hits for sequence-snippet highlighting in reports."""
    vcf_hits = list(getattr(var_info, "vcf_applied_variants", None) or [])
    snp_hits: list[dict] = []
    snp_analysis = getattr(design_summary, "snp_analysis_data", None) or {}
    if snp_analysis.get("enabled"):
        rel = getattr(var_info, "relative_pos", None)
        for hit in snp_analysis.get("hits") or []:
            if rel and _intervals_overlap(
                int(hit["template_start"]),
                int(hit["template_end"]),
                int(rel[0]),
                int(rel[1]),
            ):
                continue
            snp_hits.append(hit)
    return vcf_hits, snp_hits


def _shift_primer_pair(
    primer_pair: PrimerPairResult, display_offset: int
) -> PrimerPairResult:
    return replace(
        primer_pair,
        left_relPos_start=int(primer_pair.left_relPos_start) - display_offset,
        left_relPos_end=int(primer_pair.left_relPos_end) - display_offset,
        right_relPos_start=int(primer_pair.right_relPos_start) - display_offset,
        right_relPos_end=int(primer_pair.right_relPos_end) - display_offset,
    )


def prepare_report_sequence_view(
    var_info: AllelicVariantInfo,
    prim_settings: PrimerSettingsModel,
    primer_pair: PrimerPairResult,
    *,
    design_template: Optional[str] = None,
    allele: str,
    vcf_hits: Sequence[dict],
    snp_hits: Sequence[dict],
) -> tuple[str, str, PrimerPairResult, list[dict], list[dict], tuple[int, int]]:
    """
    Slice the primer3 design template to a report window (variant, target, both
    primers, ±REPORT_DISPLAY_FLANK). Coordinates match SEQUENCE_TEMPLATE (mutated).
    """
    full_template = (
        design_template if design_template is not None else var_info.get_seq("mutated")
    )
    orig_ref_bases = var_info.ref_bases or ""
    orig_new_bases = var_info.new_bases or ""
    var_lo, var_hi = var_info.relative_pos
    display_start, display_end = compute_report_display_bounds(
        len(full_template),
        var_lo,
        var_hi,
        prim_settings.target[0],
        prim_settings.target[1],
        primer_pair,
    )
    plain_slice = full_template[display_start:display_end]
    lo_local = var_lo - display_start
    hi_local = var_hi - display_start
    hi_local_used = hi_local
    indel_type = normalize_indel_type(var_info)
    if allele == "mut" and indel_type == IndelType.DELINS:
        new_u = (orig_new_bases or "").upper()
        if new_u:
            hi_local_used = lo_local + len(new_u) - 1
    annotated = hgvs_input_on_plain(
        plain_slice,
        lo_local,
        hi_local_used,
        indel_type,
        orig_ref_bases,
        orig_new_bases,
        allele=allele,
    )
    display_length = len(plain_slice)
    shifted_vcf = shift_template_hits_for_display(
        list(vcf_hits), display_start, display_length
    )
    shifted_snp = shift_template_hits_for_display(
        list(snp_hits), display_start, display_length
    )
    shifted_pair = _shift_primer_pair(primer_pair, display_start)
    return (
        annotated,
        plain_slice,
        shifted_pair,
        shifted_vcf,
        shifted_snp,
        (lo_local, hi_local),
    )


def visualize_sequence_as_docx(
    paragraph,
    prim_settings: PrimerSettingsModel,
    var_info: AllelicVariantInfo,
    selected_primer_pair: PrimerPairResult,
    *,
    seq_override: Optional[str] = None,
    plain_override: Optional[str] = None,
    allele: str = "wt",
    line_length: int = 60,
    vcf_hits: Optional[Sequence[dict]] = None,
    snp_hits: Optional[Sequence[dict]] = None,
    skip_snp_interval: Optional[tuple[int, int]] = None,
):
    seq = seq_override if seq_override is not None else var_info.get_seq("input")
    plain = plain_override if plain_override is not None else var_info.ref_seq

    primerF_start = int(selected_primer_pair.left_relPos_start)
    primerF_end = int(selected_primer_pair.left_relPos_end)
    primerR_start = int(selected_primer_pair.right_relPos_start)
    primerR_end = int(selected_primer_pair.right_relPos_end)

    if skip_snp_interval is None:
        rel = getattr(var_info, "relative_pos", None)
        skip_snp_interval = tuple(rel) if rel else None
    plain_highlight = build_template_highlight_lookup(
        len(plain),
        vcf_hits=vcf_hits or [],
        snp_hits=snp_hits or [],
        primer_pair=selected_primer_pair,
        skip_snp_interval=skip_snp_interval,
    )

    def _plain_highlight(idx: int) -> Optional[HighlightColor]:
        if idx < 0 or idx >= len(plain_highlight):
            return None
        return plain_highlight[idx]

    def _add_seq_run(char: str, highlight: Optional[HighlightColor]) -> None:
        nonlocal counter
        run = paragraph.add_run(char)
        run.font.name = "Courier New"
        run.font.size = Pt(11)
        _set_run_highlight(run, highlight)
        counter += 1

    # Walk annotated `seq` while mapping letters back to `plain` indices.
    plain_i = 0
    counter = 0
    ai = 0
    while ai < len(seq):
        ch = seq[ai]

        if ch != "[":
            highlight = None
            if (
                ch in "ACGTNacgtn"
                and plain_i < len(plain)
                and ch.upper() == plain[plain_i].upper()
            ):
                highlight = _plain_highlight(plain_i)
                plain_i += 1
            if highlight is None and ch == "]":
                highlight = HIGHLIGHT_VARIANT
            _add_seq_run(ch, highlight)
            if counter >= line_length:
                paragraph.add_run(f"  {format(ai+1, ',')}\n")
                counter = 0
            ai += 1
            continue

        # Bracket block
        bracket_start_plain = plain_i
        end = ai
        while end < len(seq) and seq[end] != "]":
            end += 1
        if end >= len(seq):
            end = len(seq) - 1
        inner = seq[ai + 1 : end]
        plain_from_here = plain[bracket_start_plain:]
        consumed = template_bases_consumed_by_bracket(
            inner, plain_from_here, allele=allele
        )

        # Determine if SNV bracket.
        is_snv = ">" in inner

        # For DEL/DELINS decide which side matches this allele's template.
        body = inner
        colon = body.find(":")
        if colon >= 0:
            body = body[colon + 1 :]
        ref_part = ""
        alt_part = ""
        allele_part = "ref"
        if not is_snv and "/" in body and not body.startswith("-/"):
            ref_part, alt_part = body.split("/", 1)
            alt_part = "" if alt_part == "-" else alt_part
            if alt_part and plain_from_here.upper().startswith(alt_part.upper()):
                allele_part = "alt"
            elif ref_part and plain_from_here.upper().startswith(ref_part.upper()):
                allele_part = "ref"
            elif allele == "mut":
                allele_part = "alt"

        in_bracket = False
        phase = "prefix"
        snv_phase = "prefix"
        snv_ref_seen = False
        non_snv_ref_seen = False
        bracket_plain_cursor = bracket_start_plain

        for k in range(ai, end + 1):
            c = seq[k]
            if c == "[":
                in_bracket = True
                phase = "prefix"
                snv_phase = "prefix"
                _add_seq_run(c, HIGHLIGHT_VARIANT)
                continue
            if c == "]":
                _add_seq_run(c, HIGHLIGHT_VARIANT)
                in_bracket = False
                continue
            if c == ":":
                if is_snv:
                    snv_phase = "ref"
                else:
                    phase = "ref"
                _add_seq_run(c, HIGHLIGHT_VARIANT)
                continue
            if c == ">":
                snv_phase = "alt"
                _add_seq_run(c, HIGHLIGHT_VARIANT)
                continue
            if c == "/":
                phase = "alt"
                _add_seq_run(c, HIGHLIGHT_VARIANT)
                continue

            highlight = HIGHLIGHT_VARIANT if in_bracket else None
            template_idx = None

            if c in "ACGTNacgtn":
                if is_snv:
                    # SNV without ":" (e.g. [G>A]) starts with ref base immediately.
                    if snv_phase == "prefix" and not snv_ref_seen:
                        snv_phase = "ref"
                        snv_ref_seen = True
                    # Only the allele-present base maps to the template index.
                    if (snv_phase == "ref" and allele == "wt") or (
                        snv_phase == "alt" and allele == "mut"
                    ):
                        template_idx = bracket_start_plain
                else:
                    # INS: only maps on MUT (inserted bases present in plain).
                    if body.startswith("-/"):
                        if (
                            allele == "mut"
                            and bracket_plain_cursor < bracket_start_plain + consumed
                        ):
                            template_idx = bracket_plain_cursor
                            bracket_plain_cursor += 1
                    else:
                        # DEL/DELINS without ":" starts with REF bases immediately after "["
                        if phase == "prefix" and not non_snv_ref_seen:
                            phase = "ref"
                            non_snv_ref_seen = True
                        eligible = (allele_part == "ref" and phase == "ref") or (
                            allele_part == "alt" and phase == "alt"
                        )
                        if (
                            eligible
                            and bracket_plain_cursor < bracket_start_plain + consumed
                        ):
                            template_idx = bracket_plain_cursor
                            bracket_plain_cursor += 1

            if template_idx is not None:
                region_hl = _plain_highlight(template_idx)
                if region_hl is not None:
                    highlight = region_hl
                elif in_bracket:
                    highlight = HIGHLIGHT_VARIANT

            _add_seq_run(c, highlight)

            if counter >= line_length:
                paragraph.add_run(f"  {format(k+1, ',')}\n")
                counter = 0

        plain_i = bracket_start_plain + consumed
        ai = end + 1


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    # Create the w:hyperlink tag and add needed values
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element
    new_run = OxmlElement("w:r")

    # Create a w:rPr element
    rPr = OxmlElement("w:rPr")

    # Add color if provided
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    # Underline
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)

    # Create a w:t element and set the text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _set_cell_run(
    cell, text: str, *, bold: bool = False, size_pt: Optional[float] = None
):
    """Replace cell content with a single run (for table cells)."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def add_amplicon_detail_table_to_doc(doc, primer_pair: PrimerPairResult, prim_settings):
    """
    Append a Dicey-style amplicon table when the pair has hits (insilico_status ok).
    """
    if getattr(primer_pair, "insilico_status", None) != INSILICO_OK:
        return
    amplicons = primer_pair.amplicons or []
    if not amplicons:
        return

    doc.add_heading("In-silico amplicons (Dicey)", level=2)
    ref_note = insilico_reference_description(prim_settings)
    if ref_note:
        note_p = doc.add_paragraph(ref_note)
        note_p.paragraph_format.space_after = Pt(6)

    headers = [
        "#",
        "Summary",
        "Length",
        "Penalty-Score",
        "Chrom / Gene & Transcript",
        "ForPos",
        "ForEnd",
        "RevPos",
        "RevEnd",
        "Seq (product)",
    ]
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(amplicons), cols=ncols)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, title in enumerate(headers):
        _set_cell_run(hdr_cells[j], title, bold=True, size_pt=9)

    for i, amp in enumerate(amplicons):
        row_cells = table.rows[i + 1].cells
        values = [
            str(i + 1),
            extract_amplicon_summary(amp),
            str(amp.get("Length", "") or ""),
            format_penalty_score(amp.get("Penalty")),
            amplicon_chrom_label(amp),
            str(amp.get("ForPos", "") or ""),
            str(amp.get("ForEnd", "") or ""),
            str(amp.get("RevPos", "") or ""),
            str(amp.get("RevEnd", "") or ""),
            truncate_product_seq(amp.get("Seq")),
        ]
        for j, val in enumerate(values):
            _set_cell_run(row_cells[j], val, size_pt=8)

    doc.add_paragraph()


def create_primer_report(
    designResultsSummary_obj: DesignResultsSummary, selected_primer_index: int
):
    def create_var_annotation_block(var_info: AllelicVariantInfo) -> list[str]:
        annotation = []
        if var_info.gene_ID:
            annotation += [("Gene symbol", var_info.gene_symbol)]
            annotation += [("Gene-ID", var_info.gene_ID)]
        if isinstance(var_info, TranscriptVariantInfo):
            annotation += [("Transcript-ID", var_info.transcript_id)]
        annotation += [("HGVS", create_hgvs_notation(var_info))]
        return annotation

    prim_settings = designResultsSummary_obj.primer_settings
    var_info = designResultsSummary_obj.get_variant_info()
    primer_search_results = designResultsSummary_obj.get_primer_search_results()
    allele_specific_mode = (
        isinstance(primer_search_results, dict)
        and primer_search_results.get("design_type") == "allele_specific"
    )
    if allele_specific_mode:
        wt_res = primer_search_results["wt"]
        mut_res = primer_search_results["mut"]
        common_reverse = primer_search_results.get("common_reverse_primer") or ""
        wt_pair = wt_res.primer_pairs[selected_primer_index - 1]
        mut_pair = mut_res.primer_pairs[selected_primer_index - 1]
        # Keep existing report logic intact by treating WT as "selected primer_pair"
        primer_pair = wt_pair
    else:
        primer_pair = primer_search_results.primer_pairs[selected_primer_index - 1]

    # create new document
    doc = Document()
    # Layout definition
    section = doc.sections[0]
    section.top_margin = Inches(0.7)  # Oberer Rand
    section.bottom_margin = Inches(0.7)  # Unterer Rand
    section.left_margin = Inches(0.7)  # Linker Rand
    section.right_margin = Inches(0.7)  # Rechter Rand
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)  # oder deine Wunschgröße, z. B. 12 pt

    # Add creation date to header, right-aligned
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = f"Created: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    paragraph.alignment = 2  # 2 = RIGHT

    doc.add_heading(
        (
            "Allele-specific PCR Results"
            if allele_specific_mode
            else "Primer Designer Results"
        ),
        level=0,
    )

    # add a paragraph with the variant information
    doc.add_heading("Input Variant:", level=1)
    for info, content in create_var_annotation_block(var_info):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{info}: ")
        bold_run = p.add_run(content)
        bold_run.bold = True

    primer_data = (
        ("Sequence", primer_pair.left_seq, primer_pair.right_seq),
        (
            "(Relative) start, end",
            (f"{primer_pair.left_relPos_start}, {primer_pair.left_relPos_end}"),
            (f"{primer_pair.right_relPos_start}, {primer_pair.right_relPos_end}"),
        ),
        ("Tm", f"{primer_pair.tm[0]} °C", f"{primer_pair.tm[1]} °C"),
        ("GC-content", f"{primer_pair.gc[0]} %", f"{primer_pair.gc[1]} %"),
    )

    product_data = (
        ("Product size", f"{primer_pair.product_size} bp"),
        ("Product tm", f"{primer_pair.product_tm} °C "),
    )

    from primer_designer_app.utils.primer_utils import (
        INSILICO_ERROR,
        INSILICO_NOT_APPLICABLE,
        INSILICO_OK,
        INSILICO_OK_EMPTY,
    )

    def _insilico_report_line() -> str:
        st = getattr(primer_pair, "insilico_status", None)
        if st == INSILICO_NOT_APPLICABLE:
            return "Not applicable"
        if st == INSILICO_ERROR:
            return "Error (in-silico)"
        if st == INSILICO_OK_EMPTY:
            return "0"
        if st == INSILICO_OK:
            return str(len(primer_pair.amplicons or []))
        return str(len(primer_pair.amplicons or []))

    if getattr(prim_settings, "do_insilico_pcr", False):
        if prim_settings.context == "genomic":
            product_data = product_data + (
                ("Nr of amplicons (in Genome)", _insilico_report_line()),
            )
        elif prim_settings.context == "transcriptomic":
            product_data = product_data + (
                ("Nr of amplicons (in Transcriptome)", _insilico_report_line()),
            )

    if allele_specific_mode:
        doc.add_heading("Allele reactions (AS-PCR):", level=1)
        doc.add_paragraph(
            f"Common reverse primer (shared): {common_reverse}",
            style="List Bullet",
        )
        p = doc.add_paragraph("Link to all primer-pairs: ")
        webAppHost = settings.WEB_APP_HOST  # Load WEB_APP_HOST from settings
        primer_overview_url = reverse(
            "primer_designer_app:allele_specific_primers_overview_with_uuid",
            kwargs={"uuid": designResultsSummary_obj.id},
        )
        add_hyperlink(
            p,
            f"{webAppHost}{primer_overview_url}",
            "Primer results overview",
        )
        doc.add_heading("Wild-type reaction primer pair", level=2)
    else:
        doc.add_heading("Primer Selection:", level=1)

        # Creating a table object
        p = doc.add_paragraph("Link to all primer-pairs: ")
        webAppHost = settings.WEB_APP_HOST  # Load WEB_APP_HOST from settings
        # Load primer overview URL from urls.py
        primer_overview_url = reverse(
            "primer_designer_app:snv_indel_primers_overview_with_uuid",
            kwargs={"uuid": designResultsSummary_obj.id},
        )
        add_hyperlink(
            p,
            f"{webAppHost}{primer_overview_url}",
            "Primer results overview",
        )
    table = doc.add_table(rows=1, cols=3)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[1].text = ""
    paragraph = row[1].paragraphs[0]
    run = paragraph.add_run("Forward")
    run.bold = True

    paragraph = row[2].paragraphs[0]
    run = paragraph.add_run("Reverse")
    run.bold = True

    for criteria, Info_F, Info_R in primer_data:
        row = table.add_row().cells
        row[0].text = criteria
        row[1].text = Info_F
        row[2].text = Info_R

    for criteria, info in product_data:
        doc.add_paragraph(f"{criteria}: {info}", style="List Bullet")

    add_amplicon_detail_table_to_doc(doc, primer_pair, prim_settings)

    if allele_specific_mode:
        doc.add_heading("Mutant reaction primer pair", level=2)
        primer_data_mut = (
            ("Sequence", mut_pair.left_seq, mut_pair.right_seq),
            (
                "(Relative) start, end",
                (f"{mut_pair.left_relPos_start}, {mut_pair.left_relPos_end}"),
                (f"{mut_pair.right_relPos_start}, {mut_pair.right_relPos_end}"),
            ),
            ("Tm", f"{mut_pair.tm[0]} °C", f"{mut_pair.tm[1]} °C"),
            ("GC-content", f"{mut_pair.gc[0]} %", f"{mut_pair.gc[1]} %"),
        )
        product_data_mut = (
            ("Product size", f"{mut_pair.product_size} bp"),
            ("Product tm", f"{mut_pair.product_tm} °C "),
        )
        table = doc.add_table(rows=1, cols=3)
        row = table.rows[0].cells
        row[1].text = ""
        paragraph = row[1].paragraphs[0]
        run = paragraph.add_run("Forward")
        run.bold = True
        paragraph = row[2].paragraphs[0]
        run = paragraph.add_run("Reverse")
        run.bold = True
        for criteria, Info_F, Info_R in primer_data_mut:
            row = table.add_row().cells
            row[0].text = criteria
            row[1].text = Info_F
            row[2].text = Info_R
        for criteria, info in product_data_mut:
            doc.add_paragraph(f"{criteria}: {info}", style="List Bullet")
        add_amplicon_detail_table_to_doc(doc, mut_pair, prim_settings)

    vcf_applied = getattr(var_info, "vcf_applied_variants", None) or []
    if vcf_applied:
        doc.add_heading("VCF background variants", level=1)
        doc.add_paragraph(
            f"{len(vcf_applied)} variant(s) from the uploaded VCF were spiked into "
            "the reference sequence before primer design."
        )
        table = doc.add_table(rows=1 + len(vcf_applied), cols=3)
        table.style = "Table Grid"
        for j, title in enumerate(["ID", "Genomic", "REF → ALT"]):
            _set_cell_run(table.rows[0].cells[j], title, bold=True, size_pt=9)
        for i, row in enumerate(vcf_applied):
            cells = table.rows[i + 1].cells
            values = [
                row.get("id", ""),
                f"chr{row.get('chrom', '')}:{row.get('pos', '')}",
                f"{row.get('ref', '')} → {row.get('alt', '')}",
            ]
            for j, val in enumerate(values):
                _set_cell_run(cells[j], str(val), size_pt=9)

    snp_analysis = getattr(designResultsSummary_obj, "snp_analysis_data", None) or {}
    if snp_analysis.get("enabled"):
        doc.add_heading("Known SNPs (gnomAD, MAF > 1%)", level=1)
        doc.add_paragraph(snp_analysis.get("message", ""))
        if snp_analysis.get("region"):
            region = snp_analysis["region"]
            doc.add_paragraph(
                f"Design region: chr{region.get('chromosome')}:"
                f"{region.get('start')}-{region.get('end')}",
                style="List Bullet",
            )
        pair_status = getattr(primer_pair, "snp_status", None)
        if pair_status:
            doc.add_paragraph(
                f"Selected primer pair SNP risk: {pair_status}",
                style="List Bullet",
            )
        conflicts = getattr(primer_pair, "snp_conflicts", None) or []
        if conflicts:
            doc.add_heading("SNPs contained in PCR amplified region", level=2)
            table = doc.add_table(rows=1 + len(conflicts), cols=5)
            table.style = "Table Grid"
            for j, title in enumerate(["ID", "Genomic", "Alleles", "MAF", "Primer"]):
                _set_cell_run(table.rows[0].cells[j], title, bold=True, size_pt=9)
            for i, hit in enumerate(conflicts):
                row_cells = table.rows[i + 1].cells
                maf_val = hit.get("maf")
                maf_str = f"{maf_val:.4f}" if maf_val is not None else ""
                values = [
                    hit.get("id", ""),
                    f"{hit.get('genomic_start')}–{hit.get('genomic_end')}",
                    hit.get("alleles", ""),
                    maf_str,
                    hit.get("primer", "-- no primer overlap --"),
                ]
                for j, val in enumerate(values):
                    _set_cell_run(row_cells[j], val, size_pt=8)

    # 3. Sequence visualization
    doc.add_heading("Sequence snippet:", level=1)
    doc.add_paragraph(
        "Sequence snippet shows the region of interest (input variant, target region, "
        f"selected primers, ±{REPORT_DISPLAY_FLANK} bp flank). UTRs and introns are "
        "lowercase; exons are uppercase."
    )

    doc.add_heading("Legend:", level=2)
    vcf_hits, snp_hits = _report_region_hits(designResultsSummary_obj, var_info)
    legend_content = [
        ["Primer", HIGHLIGHT_PRIMER],
        ["Input variant", HIGHLIGHT_VARIANT],
    ]
    if vcf_hits:
        legend_content.append(["VCF spiked variant", HIGHLIGHT_VCF])
    if snp_hits:
        legend_content.append(["Common SNP (gnomAD)", HIGHLIGHT_SNP])
        legend_content.append(["SNP at primer binding site", HIGHLIGHT_SNP_CONFLICT])

    for region, color in legend_content:
        list_obj = doc.add_paragraph(style="List Bullet")
        run = list_obj.add_run(region)
        _set_run_highlight(run, color)

    if allele_specific_mode:
        doc.add_heading("Wild-type reaction", level=2)
        wt_annotated, wt_slice, wt_shifted, _, _, wt_skip = (
            prepare_report_sequence_view(
                var_info,
                prim_settings,
                wt_pair,
                design_template=var_info.ref_seq,
                allele="wt",
                vcf_hits=[],
                snp_hits=[],
            )
        )
        paragraph = doc.add_paragraph()
        visualize_sequence_as_docx(
            paragraph,
            prim_settings,
            var_info,
            wt_shifted,
            seq_override=wt_annotated,
            plain_override=wt_slice,
            allele="wt",
            vcf_hits=[],
            snp_hits=[],
            skip_snp_interval=wt_skip,
        )

        (
            mut_annotated,
            mut_slice,
            mut_shifted,
            mut_vcf,
            mut_snp,
            mut_skip,
        ) = prepare_report_sequence_view(
            var_info,
            prim_settings,
            mut_pair,
            design_template=var_info.get_seq("mutated"),
            allele="mut",
            vcf_hits=vcf_hits,
            snp_hits=snp_hits,
        )

        doc.add_heading("Mutant reaction", level=2)
        paragraph = doc.add_paragraph()
        visualize_sequence_as_docx(
            paragraph,
            prim_settings,
            var_info,
            mut_shifted,
            seq_override=mut_annotated,
            plain_override=mut_slice,
            allele="mut",
            vcf_hits=mut_vcf,
            snp_hits=mut_snp,
            skip_snp_interval=mut_skip,
        )
    else:
        (
            seq_annotated,
            plain_slice,
            shifted_pair,
            shifted_vcf,
            shifted_snp,
            skip_interval,
        ) = prepare_report_sequence_view(
            var_info,
            prim_settings,
            primer_pair,
            allele="mut",
            vcf_hits=vcf_hits,
            snp_hits=snp_hits,
        )
        paragraph = doc.add_paragraph()
        visualize_sequence_as_docx(
            paragraph,
            prim_settings,
            var_info,
            shifted_pair,
            seq_override=seq_annotated,
            plain_override=plain_slice,
            allele="mut",
            vcf_hits=shifted_vcf,
            snp_hits=shifted_snp,
            skip_snp_interval=skip_interval,
        )

    # Save the document to an in-memory file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _init_report_document(title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = f"Created: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    paragraph.alignment = 2

    doc.add_heading(title, level=0)
    return doc


def _add_bullet_field(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{label}: ")
    bold_run = p.add_run(value)
    bold_run.bold = True


def _add_sv_primer_pairs_table(doc: Document, primer_rows: list) -> None:
    headers = [
        "Rank",
        "Forward primer",
        "Reverse primer",
        "Penalty",
        "Product size (bp)",
        "Primer Tm (°C)",
        "Primer GC (%)",
        "Forward genomic",
        "Reverse genomic",
    ]
    if not primer_rows:
        doc.add_paragraph("No primer pairs found for this window.")
        return

    table = doc.add_table(rows=1 + len(primer_rows), cols=len(headers))
    table.style = "Table Grid"
    for j, title in enumerate(headers):
        _set_cell_run(table.rows[0].cells[j], title, bold=True, size_pt=9)

    for i, row in enumerate(primer_rows):
        pair = row["pair"]
        genomic = row["genomic_positions"]
        tm_text = (
            f"{pair.tm[0]}, {pair.tm[1]}" if pair.tm and len(pair.tm) >= 2 else "—"
        )
        gc_text = (
            f"{pair.gc[0]}, {pair.gc[1]}" if pair.gc and len(pair.gc) >= 2 else "—"
        )
        values = [
            str(i + 1),
            pair.left_seq,
            pair.right_seq,
            str(pair.penalty),
            str(pair.product_size),
            tm_text,
            gc_text,
            f"{genomic['forward_start']}–{genomic['forward_end']}",
            f"{genomic['reverse_start']}–{genomic['reverse_end']}",
        ]
        for j, val in enumerate(values):
            _set_cell_run(table.rows[i + 1].cells[j], val, size_pt=8)


def create_structural_variant_primer_report(
    design_results_summary: DesignResultsSummary,
) -> io.BytesIO:
    """
    Word report for structural variant primer design: all primer pairs per window,
    without sequence visualization or primer-pair selection.
    """
    sv_info = design_results_summary.get_structural_variant_info_data()
    sv_results = design_results_summary.get_sv_primer_results()
    prim_settings = design_results_summary.primer_settings

    doc = _init_report_document("Structural Variant Primer Designer Results")

    doc.add_heading("Structural variant query", level=1)
    chromosome = sv_info.get("chromosome", "")
    start_pos = sv_info.get("start_position", "")
    end_pos = sv_info.get("end_position", "")
    _add_bullet_field(doc, "Reference genome", sv_info.get("reference_genome", ""))
    _add_bullet_field(doc, "Chromosome", f"chr{chromosome}")
    _add_bullet_field(doc, "Start position", str(start_pos))
    _add_bullet_field(doc, "End position", str(end_pos))
    if start_pos and end_pos:
        span = int(end_pos) - int(start_pos) + 1
        _add_bullet_field(doc, "Span", f"{span} bp")

    doc.add_heading("Design windows", level=2)
    for window in sv_info.get("windows", []):
        doc.add_paragraph(
            f"{window.get('label', '').replace('_', ' ').title()}: "
            f"{window.get('window_start_genomic')}–{window.get('window_end_genomic')} "
            f"(chr{chromosome})",
            style="List Bullet",
        )

    doc.add_heading("Primer design parameters", level=1)
    if prim_settings:
        _add_bullet_field(doc, "Optimal Tm", f"{prim_settings.tm} °C")
        _add_bullet_field(doc, "GC target", f"{prim_settings.gc} %")
        _add_bullet_field(doc, "Max poly-X", str(prim_settings.max_poly_x))
        if prim_settings.productsize_range:
            prod_range = (
                f"{prim_settings.productsize_range[0]}–"
                f"{prim_settings.productsize_range[1]} bp"
            )
            _add_bullet_field(
                doc,
                "Product size range",
                prod_range,
            )

    doc.add_heading("Designed primer pairs", level=1)
    doc.add_paragraph(
        "All primer pairs returned by Primer3 for each design window are listed below. "
        "Genomic coordinates refer to the selected reference assembly."
    )

    labels_in_report = [label for label in SV_WINDOW_ORDER if label in sv_results]
    labels_in_report.extend(
        label for label in sv_results if label not in labels_in_report
    )

    for label in labels_in_report:
        window_result = sv_results[label]
        window = window_result["window"]
        window_title = label.replace("_", " ").title()
        doc.add_heading(window_title, level=2)
        doc.add_paragraph(
            f"Window coordinates (chr{chromosome}): "
            f"{window['window_start_genomic']}–{window['window_end_genomic']}"
        )
        _add_sv_primer_pairs_table(doc, window_result.get("primer_rows", []))
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
